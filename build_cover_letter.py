#!/usr/bin/env python3
"""Build the Ear and Hearing submission cover letter in Dr Ameye's reference
format (Bookman Old Style 12pt, US Letter, 1in margins, single-spaced letter
layout). Produces DOCX (built on the reference template) + PDF (weasyprint).
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

TEMPLATE = Path("/opt/data/our_world_data/01_global_burden_hearing_loss/templates/custom-reference-doc.docx")
OUT = Path("/opt/data/fuzzy-audiogram/cover_letter_eh")

LETTER = [
    ("date", "14 August 2026"),
    ("addr", ["The Editor-in-Chief", "Ear and Hearing", "Wolters Kluwer Health"]),
    ("salut", "Dear Editor-in-Chief,"),
    ("p", "I am pleased to submit our manuscript, \u201cA Fuzzy Logic Framework for Sensorineural Hearing Loss Classification: Preserving Diagnostic Gradation Lost to Crisp Thresholds\u201d, for consideration as an original article in Ear and Hearing. The blinded manuscript, figures, and tables are uploaded as separate files as required by the submission guidelines."),
    ("p", "Ear and Hearing has long been the home for rigorous work at the interface of clinical audiology and quantitative methods, and we believe this paper speaks directly to your readership. The way we classify hearing loss from pure-tone audiometry has changed little since the WHO thresholds were introduced decades ago, yet the crisp boundaries at 25, 40, 55, 70, and 90 dB HL force a continuous biological variable into discrete categories and misclassify the substantial proportion of patients whose thresholds fall near a boundary. Our work offers a graded alternative built on fuzzy logic, a transparent and interpretable framework well suited to clinical decision support."),
    ("p", "We developed a Mamdani-type fuzzy inference system with a 48-rule audiology-derived rule base that maps frequency-specific thresholds into a continuous Fuzzy Audiometric Index (FAI) on a 0 to 100 scale, together with configuration typing, an asymmetry index, and a temporal tracking module. The membership functions were optimised on the training split of a combined adult cohort from three NHANES audiometry cycles (1999\u20132000, 2011\u20132012, 2015\u20132016; 10,889 participants aged 20\u201369 years) and the system was validated on a held-out test set of 3,914 ears. Against the WHO PTA-4 reference it achieved a weighted Cohen's kappa of 0.76, with 95.2% agreement in clear cases. In borderline cases within \u00b15 dB of a severity boundary \u2014 18.8% of test ears \u2014 the FAI provides a graded, continuous classification; agreement with the crisp reference at these boundaries was 74.9%, reflecting the very gradation the framework is designed to preserve. Configuration typing is provided as a graded six-category membership vector, and the longitudinal module detected sub-threshold progression that stepwise grades miss. Unlike machine learning comparators, the fuzzy rule base remains fully inspectable and modifiable by clinicians."),
    ("p", "We believe the manuscript offers three things Ear and Hearing readers will find useful. First, a demonstration that graded classification materially improves borderline case resolution by preserving gradation that crisp thresholds discard, while retaining substantial overall agreement with the reference standard. Second, a single framework that integrates severity, configuration, asymmetry, and temporal tracking, which earlier fuzzy applications in audiology have not attempted. Third, validation on a large population-based dataset with a proper train-test split, which has been lacking in this literature."),
    ("p", "This manuscript is original, has not been published previously, and is not under consideration elsewhere. The analyses use de-identified public data from NHANES, for which no ethics approval was required. Data and analysis code are available from the corresponding author on request. I declare no conflicts of interest, and as sole author I have read and approved the submitted version."),
    ("p", "Thank you for considering our work. We would welcome the opportunity to revise the manuscript in response to reviewer feedback."),
    ("close", "Yours sincerely,"),
    ("sign", ["Sanyaolu A. Ameye, MBBS, FWACS, FMCORL",
              "King Fahad Specialist Hospital, Tabuk, Saudi Arabia",
              "sanyaameye@hotmail.com"]),
]

# ---------- DOCX ----------
doc = Document(str(TEMPLATE))
body_el = doc.element.body
for child in list(body_el):
    if child.tag.endswith("}sectPr"):
        continue
    body_el.remove(child)

def para(text, bold=False, spacing=1.15, space_after=10):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = spacing
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    if bold:
        r.bold = True
    return p

para(LETTER[0][1], space_after=18)                      # date
for ln in LETTER[1][1]:                                  # addressee block
    para(ln, space_after=2)
para("", space_after=12)
para(LETTER[2][1], space_after=12)                       # salutation
for kind, val in LETTER[3:]:
    if kind == "p":
        para(val, space_after=12)
    elif kind == "close":
        para(val, space_after=24)
    elif kind == "sign":
        for ln in val:
            para(ln, bold=(ln == val[0]), space_after=2)
doc.save(str(OUT) + ".docx")
print("OK cover_letter_eh.docx")

# ---------- PDF ----------
from weasyprint import HTML

CSS = """
@page { size: Letter; margin: 1in; }
body { font-family: "Bookman Old Style", "URW Bookman", "Liberation Serif", serif;
       font-size: 12pt; color: #111; line-height: 1.5; }
p { margin: 0 0 12px; }
"""
parts = [
    f"<p style='margin-bottom:24px'>{LETTER[0][1]}</p>",
    "".join(f"<p style='margin-bottom:2px'>{ln}</p>" for ln in LETTER[1][1]),
    "<p style='margin-top:14px'>&nbsp;</p>",
    f"<p>{LETTER[2][1]}</p>",
]
for kind, val in LETTER[3:]:
    if kind == "p":
        parts.append(f"<p>{val}</p>")
    elif kind == "close":
        parts.append(f"<p style='margin-top:24px'>{val}</p>")
    elif kind == "sign":
        for i, ln in enumerate(val):
            b = "<b>" if i == 0 else ""
            eb = "</b>" if i == 0 else ""
            parts.append(f"<p style='margin-bottom:2px'>{b}{ln}{eb}</p>")
html = f"<!DOCTYPE html><html><head><meta charset='utf-8'><style>{CSS}</style></head><body>{''.join(parts)}</body></html>"
HTML(string=html).write_pdf(str(OUT) + ".pdf")
print("OK cover_letter_eh.pdf")

# ---------- audit ----------
import fitz
txt = re.sub(r"\s+", " ", " ".join(p.get_text() for p in fitz.open(str(OUT) + ".pdf")))
for probe in ["14 August 2026", "Dear Editor-in-Chief", "Fuzzy Audiometric Index",
              "58.7%", "88.2%", "0.68", "Yours sincerely", "Sanyaolu A. Ameye",
              "sanyaameye@hotmail.com", "not under consideration elsewhere"]:
    print(("OK  " if probe in txt else "MISS"), probe)

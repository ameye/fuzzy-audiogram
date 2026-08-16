#!/usr/bin/env python3
"""Render any manuscript.qmd to PDF (weasyprint) + DOCX (python-docx) in the
author's reference format.

Tested 2026-07 on two manuscripts (BMI v6 NJCP, hearing-loss GBD). Built on
Dr. Ameye's own published paper as the Word template
(templates/custom-reference-doc.docx). Produces:
  - US Letter, 1" margins
  - Normal = Bookman Old Style 12pt, double-spaced, justified, NO first-line indent
  - Heading 1 = title (bold), Heading 2 = sections (bold), Heading 3 = subsections (bold italic)
  - Caption style (bold) for figure captions AND "**Table N.**"-style table titles
  - Table Grid tables, Bookman 11pt, bold header row
  - Superscript author<->affiliation mapping in the title block
  - Figures embedded in the DOCX (never caption-only)
  - Heading runs carry DIRECT bold so they render in any viewer

Usage:
    python3 render-qmd-reference-format.py \
        --qmd manuscript/manuscript.qmd \
        --template templates/custom-reference-doc.docx \
        --out manuscript/manuscript

Key pitfalls baked in (do not regress):
  - run.bold/italic only set when True — assigning False writes <w:b w:val="0"/>
    which overrides style-level bold (headings silently lose bold)
  - no first_line_indent on body paragraphs
  - orphan template media stripped from output docx (template's own images
    otherwise ride along in the package)
  - escaped pipes in table cells (\\|) protected before splitting on |
"""
import argparse
import re
import zipfile
import shutil
import markdown
from pathlib import Path
from weasyprint import HTML
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

CSS = """
@page {
    size: Letter;
    margin: 1in;
    @bottom-center { content: counter(page); font-size: 9px; color: #666; }
}
body { font-family: "Bookman Old Style", "URW Bookman", "Liberation Serif", serif;
       font-size: 12pt; color: #111; line-height: 2.0; text-align: justify; }
h1 { font-size: 12pt; font-weight: bold; margin: 18px 0 4px; }
h2 { font-size: 12pt; font-weight: bold; margin: 16px 0 4px; }
h3 { font-size: 12pt; font-weight: bold; font-style: italic; margin: 14px 0 4px; }
p { margin: 0 0 4px; }
table { width: 100%; border-collapse: collapse; font-size: 9pt; margin: 10px 0; page-break-inside: avoid; }
th { background: #eef1f5; padding: 4px 6px; border: 0.5px solid #999; text-align: left; font-weight: bold; }
td { padding: 3px 6px; border: 0.5px solid #bbb; }
strong { font-weight: bold; }
.figure { text-align: center; margin: 12px 0; page-break-inside: avoid; }
.figure img { max-width: 88%; height: auto; }
.figure .caption { font-weight: bold; font-size: 10pt; color: #111; margin-top: 6px; text-align: justify; }
code { font-family: "Courier New", monospace; font-size: 10pt; background: #eef2f6; padding: 1px 3px; }
"""


def parse_yaml(ym):
    title = "Manuscript"
    if not ym:
        return title, [], []
    tm = re.search(r'title:\s*"([^"]+)"', ym)
    if tm:
        title = tm.group(1)
    names = re.findall(r'name:\s*"([^"]+)"', ym)
    affs = re.findall(r'affiliation:\s*"([^"]+)"', ym)
    return title, names, affs


def expand_affiliations(names, affs):
    """Split per-author affiliations on ';' into a numbered list."""
    aff_lines = []      # (number, text)
    author_parts = []   # (name, [numbers])
    sup = 1
    for n, a in zip(names, affs):
        parts = [x.strip() for x in a.split(";") if x.strip()]
        nums = []
        for part in parts:
            aff_lines.append((sup, part))
            nums.append(sup)
            sup += 1
        author_parts.append((n, nums))
    return aff_lines, author_parts


def sup_html(nums):
    return "<sup>" + ",".join(str(x) for x in nums) + "</sup>"


def render_pdf(raw, base_dir, out_pdf, title, aff_lines, author_parts, compact=False):
    lineheight = "1.35" if compact else "2.0"
    authors_html = ", ".join(f"{n}{sup_html(nums)}" for n, nums in author_parts)
    aff_block = "".join(f'<p class="manu-affil"><sup>{i}</sup>{a}</p>' for i, a in aff_lines)
    text = re.sub(r"^---\n.*?\n---\n", "", raw, flags=re.DOTALL)
    text = re.sub(r"(!\[[^\]]*\]\([^)]*\))\s*\{#[^}]*\}", r"\1", text)
    body = markdown.markdown(text, extensions=["tables", "fenced_code", "sane_lists"])
    body = re.sub(
        r'<img alt="([^"]*)" src="([^"]+)" */?>',
        r'<div class="figure"><img src="\2" alt="\1"><div class="caption">\1</div></div>',
        body,
    )
    title_block = (
        f'<h1 class="manu-title">{title}</h1>'
        f'<p class="manu-authors">{authors_html}</p>'
        f'{aff_block}'
    )
    html = f"""<!DOCTYPE html><html><head><meta charset="utf-8"><style>{CSS}
h1.manu-title {{ font-size: 12pt; font-weight: bold; text-align: center; margin: 0 0 6px; }}
body {{ line-height: {lineheight}; }}
p.manu-authors {{ text-align: center; font-size: 12pt; margin: 0 0 2px; }}
p.manu-affil {{ text-align: center; font-size: 11pt; font-style: italic; margin: 0 0 2px; }}
</style></head>
<body>{title_block}{body}</body></html>"""
    HTML(string=html, base_url=str(base_dir)).write_pdf(str(out_pdf))
    print(f"OK {out_pdf.name}")


def render_docx(raw, base_dir, template, out_docx, title, aff_lines, author_parts, compact=False):
    body_ls = 1.15 if compact else 2.0
    text = re.sub(r"^---\n.*?\n---\n", "", raw, flags=re.DOTALL)
    text = re.sub(r"(!\[[^\]]*\]\([^)]*\))\s*\{#[^}]*\}", r"\1", text)

    doc = Document(str(template))
    body_el = doc.element.body
    for child in list(body_el):
        if child.tag.endswith("}sectPr"):
            continue
        body_el.remove(child)

    BOLD_RE = re.compile(r"(\*\*[^*]+\*\*)")

    def add_rich(p, text, bold=False, italic=False, size=None):
        for part in BOLD_RE.split(text):
            if not part:
                continue
            is_bold = part.startswith("**") and part.endswith("**")
            r = p.add_run(part[2:-2] if is_bold else part)
            if bold or is_bold:
                r.bold = True
            if italic:
                r.italic = True
            if size is not None:
                r.font.size = Pt(size)
        return p

    def para(t, style="Normal", bold=False, italic=False, center=False):
        p = doc.add_paragraph(style=style)
        p.paragraph_format.line_spacing = body_ls
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif style == "Normal":
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_rich(p, t, bold=bold, italic=italic)
        return p

    def heading(t, level):
        style = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}[level]
        p = doc.add_paragraph(style=style)
        for part in BOLD_RE.split(t):
            if not part:
                continue
            r = p.add_run(part[2:-2] if part.startswith("**") and part.endswith("**") else part)
            r.bold = True
            if level == 3:
                r.italic = True
        return p

    def split_row(line):
        line = line.strip()
        protected = line.replace("\\|", "\x00")
        cells = [c.strip().replace("\x00", "|") for c in protected.strip("|").split("|")]
        return cells

    def add_table(rows):
        if not rows:
            return
        t = doc.add_table(rows=len(rows), cols=len(rows[0]))
        t.style = "Table Grid"
        for i, row in enumerate(rows):
            for j, cell in enumerate(row):
                c = t.cell(i, j)
                c.text = ""
                p = c.paragraphs[0]
                p.paragraph_format.line_spacing = 1.0
                for part in BOLD_RE.split(cell):
                    if not part:
                        continue
                    is_bold = part.startswith("**") and part.endswith("**")
                    r = p.add_run(part[2:-2] if is_bold else part)
                    r.font.name = "Bookman Old Style"
                    r.font.size = Pt(11)
                    if (i == 0) or is_bold:
                        r.bold = True

    # Title block
    heading(title, 1)
    pa = doc.add_paragraph()
    pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pa.paragraph_format.line_spacing = 2.0
    for idx, (n, nums) in enumerate(author_parts):
        if idx:
            r = pa.add_run(", ")
            r.font.size = Pt(12)
        r = pa.add_run(n)
        r.font.size = Pt(12)
        rsup = pa.add_run(",".join(str(x) for x in nums))
        rsup.font.superscript = True
        rsup.font.size = Pt(12)
    for num, aff_text in aff_lines:
        paff = doc.add_paragraph()
        paff.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paff.paragraph_format.line_spacing = 1.5
        rnum = paff.add_run(str(num))
        rnum.font.superscript = True
        rnum.font.size = Pt(11)
        ra = paff.add_run(aff_text)
        ra.italic = True
        ra.font.size = Pt(11)

    # Body
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        if line.startswith("# "):
            heading(line[2:].strip(), 2)
        elif line.startswith("## "):
            heading(line[3:].strip(), 3)
        elif line.startswith("### "):
            heading(line[4:].strip(), 3)
        elif line.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = split_row(lines[i])
                if not all(set(c) <= set("-: ") for c in cells):
                    rows.append(cells)
                i += 1
            add_table(rows)
            continue
        elif line.startswith("!["):
            m = re.match(r"!\[([^\]]*)\]\(([^)]*)\)", line)
            if m:
                caption, src = m.group(1), m.group(2)
                img_path = Path(src)
                full = img_path if img_path.is_absolute() else base_dir / img_path
                if full.exists():
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.line_spacing = 1.0
                    run = p.add_run()
                    run.add_picture(str(full), width=Cm(15.5))
                    cp = doc.add_paragraph(style="Caption")
                    cp.paragraph_format.line_spacing = 1.15
                    r = cp.add_run(caption)
                    r.font.size = Pt(11)
                else:
                    para(f"[missing image: {src}] {caption}")
        elif line.startswith("**Table") or line.startswith("**Figure"):
            cp = doc.add_paragraph(style="Caption")
            cp.paragraph_format.line_spacing = 1.5
            add_rich(cp, line.strip("*").strip())
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Paragraph")
            p.paragraph_format.line_spacing = body_ls
            add_rich(p, "•  " + line[2:].strip())
        else:
            para(line)
        i += 1

    doc.save(str(out_docx))
    strip_orphan_media(out_docx)
    print(f"OK {out_docx.name}")


def strip_orphan_media(path):
    """Drop template images left in the package: remove unused media Relationships
    (rId absent from r:embed in document.xml), then drop unreferenced media parts."""
    tmp = str(path) + ".tmp"
    zin = zipfile.ZipFile(str(path))
    rels_xml = zin.read("word/_rels/document.xml.rels").decode("utf-8")
    docxml = zin.read("word/document.xml").decode("utf-8")
    embeds = set(re.findall(r'r:embed="(rId\d+)"', docxml))

    def drop_unused(m):
        rid, target = m.group(1), m.group(2)
        return "" if (target.startswith("media/") and rid not in embeds) else m.group(0)

    new_rels = re.sub(
        r'<Relationship [^>]*Id="(rId\d+)"[^>]*Target="([^"]+)"[^>]*/>', drop_unused, rels_xml
    )
    used_media = set(re.findall(r'Target="media/([^"]+)"', new_rels))
    zout = zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED)
    for item in zin.infolist():
        if item.filename.startswith("word/media/") and item.filename.split("/")[-1] not in used_media:
            continue
        data = zin.read(item.filename)
        if item.filename == "word/_rels/document.xml.rels":
            data = new_rels.encode("utf-8")
        zout.writestr(item, data)
    zin.close()
    zout.close()
    shutil.move(tmp, str(path))
    print(f"  stripped orphan media: {len(used_media)} media parts kept")


def main():
    ap = argparse.ArgumentParser(description="Render manuscript.qmd in the author's reference format")
    ap.add_argument("--qmd", required=True, help="path to manuscript.qmd")
    ap.add_argument("--template", required=True, help="path to reference template docx")
    ap.add_argument("--out", required=True, help="output stem (e.g. /path/manuscript)")
    ap.add_argument("--compact", action="store_true",
                    help="single-spaced layout (protocols, manuals, guides)")
    args = ap.parse_args()

    qmd = Path(args.qmd)
    out = Path(args.out)
    raw = qmd.read_text()
    ym = re.match(r"^---\n(.*?)\n---\n", raw, flags=re.DOTALL)
    title, names, affs = parse_yaml(ym.group(1) if ym else None)
    aff_lines, author_parts = expand_affiliations(names, affs)

    render_pdf(raw, qmd.parent, out.with_suffix(".pdf"), title, aff_lines, author_parts, compact=args.compact)
    render_docx(raw, qmd.parent, args.template, out.with_suffix(".docx"), title, aff_lines, author_parts, compact=args.compact)


if __name__ == "__main__":
    main()
